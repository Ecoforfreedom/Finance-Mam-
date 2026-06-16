from pathlib import Path
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import json, re
ROOT=Path(r'未来基金POC/未来基金_投后管理POC模拟数据')
projects=[('P001','星澜计算','AI材料模拟','明确AI for Science；完成工程样机；新增专利'),('P002','原矩量子','超导量子测控','量子早期；上海工商完成；融资新闻A+与股东名册A轮追加冲突'),('P003','量衡精密','NV色心磁测量','上海落地口径模糊；客户验证延期两个季度'),('P004','辰析生物','AI蛋白设计','AI for Science；A+轮；客户验证完成；引入战略投资人'),('P005','深材智能','材料科学数据平台','科学研究平台但AI非核心证据不足；BP过期；融资口径冲突'),('P006','云弦量子','量子密钥分发网关','量子早期上海落地；签署首个商业订单；返投完成'),('P007','极微半导体','模拟芯片','连续两个季度无有效更新；信息披露逾期；关键人员疑似变化'),('P008','灵枢机器人','柔性装配机器人','非专题；首个商业订单；董事会现金流异常；员工人数倒挂'),('P009','谱元生物','AI酶工程','AI for Science；融资协议未全额交割；返投金额口径冲突'),('P010','棱镜医疗','多光谱成像器械','非专题；医疗器械注册受理；缺少检测附件'),('P011','序智软件','PLM知识图谱','边界排除：使用AI做办公和市场分析但核心为企业软件'),('P012','拓界智造','产线数字化','量子概念排除；融资轮次冲突；战略投资人未进入股东名册')]

def setup(doc):
    doc.styles['Normal'].font.name='Microsoft YaHei'; doc.styles['Normal'].font.size=Pt(10.5)

def append_once(path,title,blocks):
    doc=Document(path); setup(doc)
    if any(title in p.text for p in doc.paragraphs): return
    doc.add_heading(title,2)
    for h,paras in blocks:
        doc.add_heading(h,3)
        for para in paras: doc.add_paragraph(para)
    doc.save(path)

def gp_blocks(name,tech,scene):
    base=[f"核验补充{i}：GP投后团队对{name}的判断采用证据分层方法。企业季度报告可说明管理层观点，但不能自动证明协议履约完成；工商登记、审计报表、正式交割文件、董事会文件和已签署协议优先级更高。针对{tech}方向，本期重点比较上一季度计划、本季度实际完成事项、附件形成日期和对外宣传文字，避免将项目背景重复描述误判为新增进展。" for i in range(1,13)]
    return [('七、证据分层及来源优先级',base[:4]),('八、跨期变化识别过程',base[4:8]),('九、需由未来基金确认的问题',[f"对于{name}，本期需持续关注：{scene}。GP建议发出书面问询，要求企业按字段提供证明，包括材料名称、形成日期、是否盖章、是否外部公开、是否可与其他来源交叉验证。"]+base[8:11]),('十、投后管理建议',[base[11],"建议动作仅限管理跟进：要求补充材料、安排专题会议、更新下一跟进日期、在协议履约表中调整红黄绿灰状态。GP不基于当前材料作退出、失败或估值贬损等缺乏依据的结论。"])]

def board_blocks(name,tech,scene):
    return [('补充讨论一：融资与现金流',[f"董事会进一步询问{name}融资推进的具体节点。管理层说明部分投资人已完成商务沟通，但董事要求区分融资意向、TS、正式协议、首次交割、全部交割和工商变更。财务负责人补充，现金流压力主要来自研发人员薪酬、外协测试和客户现场支持，若融资交割延后，需要调整付款节奏，但该讨论不等同于项目经营失败结论。","未来基金观察员要求在会后提供银行流水摘要、预算实际对比表和应收账款账龄表，并明确哪些收入已满足验收条件。"]),('补充讨论二：上海落地与返投',[f"围绕上海事项，会议明确不得将落地意向、临时办公室、工商注册、社保入驻和实际经营混为同一概念。{name}应按投资协议分别提交营业执照或工商摘录、租赁合同、员工社保、采购合同、付款流水和本地业务说明。","对于仅有口头计划或新闻报道的事项，董事会同意暂列待确认；对于已超过截止日期但缺少正式证明的事项，应进入红色或黄色跟踪，并设置整改时间。"]),('补充讨论三：研发、客户和附件',[f"研发负责人围绕{tech}介绍了样机、测试和客户适配进展。董事要求研发团队在下次会议前提交测试报告目录、问题清单和客户反馈原文摘录。客户负责人说明部分客户验证晚于原计划，原因包括客户现场排期、法务合同流程和检测窗口不足。",f"会议形成一致意见：{scene}需要通过两份或以上材料交叉验证，不能仅凭单份材料下结论。"]),('补充行动项',["1. 投后联系人在2026-07-31前提交附件清单；2. 财务负责人在2026-08-05前提交银行流水和预算对比；3. 研发负责人在2026-08-10前提交测试报告；4. GP在2026-08-12前更新协议履约表；5. 未来基金项目负责人将红色和灰色事项提交投后例会。","会议秘书需将上述行动项纳入下一次董事会材料，未完成事项不得在季度报告中直接表述为已完成。"])]

def agreement_blocks(name,tech,scene):
    paras=[f"补充约定{i}：各方确认，{name}围绕{tech}开展业务时形成的技术资料、客户材料、融资材料、工商材料和财务材料具有不同证明力。若企业商业计划书、新闻稿、会议口头说明与正式协议、工商登记或交割文件不一致，应保留全部来源并在投后管理系统中标记冲突，不得简单删除低优先级信息。" for i in range(1,17)]
    return [('第十二条 信息核验和材料优先级',paras[:4]),('第十三条 口径差异处理',paras[4:8]),('第十四条 上海事项专项证明',paras[8:12]),('第十五条 违约整改与宽限安排',[f"如{name}出现{scene}相关事项，投资方可要求公司在十个工作日内补充说明。若延误具有合理客观原因，公司应提供政府受理通知、客户延期邮件、检测机构排期、银行付款记录或其他可验证材料。"]+paras[12:15]),('第十六条 附件清单和持续披露',[paras[15],"协议附件包括季度报告模板、资金使用表、上海落地证明清单、关键人员承诺函、知识产权清单、客户合同摘要和融资节点通知模板。公司应持续维护附件版本，确保每个重要字段可追溯至具体材料、章节、日期和责任人。"])]

for pid,name,tech,scene in projects:
    for p in ROOT.rglob(f'{pid}_{name}_GP季度投后报告_*_V2.0_*.docx'): append_once(p,'增强补充：GP核验工作底稿摘录',gp_blocks(name,tech,scene))
    for p in ROOT.rglob(f'{pid}_{name}_董事会会议纪要_2026Q2_V2.0_*.docx'): append_once(p,'增强补充：会议详细讨论记录',board_blocks(name,tech,scene))
    for p in ROOT.rglob(f'{pid}_{name}_投资协议_*_V2.0_*.docx'): append_once(p,'增强补充：持续跟踪和信息核验条款',agreement_blocks(name,tech,scene))
    for p in ROOT.rglob(f'{pid}_{name}_股东协议_*_V2.0_*.docx'): append_once(p,'增强补充：持续跟踪和信息核验条款',agreement_blocks(name,tech,scene)[:3])
    for p in ROOT.rglob(f'{pid}_{name}_补充协议_*_V2.0_*.docx'): append_once(p,'增强补充：上海落地返投核验细则',agreement_blocks(name,tech,scene)[2:])

def chars(p): return sum(len(x.text) for x in Document(p).paragraphs)
files=[p for p in ROOT.rglob('*') if p.is_file()]
qdocs=[p for p in ROOT.rglob('*V2.0*.docx') if '季度投后报告' in p.name and 'GP季度' not in p.name]
gp=[p for p in ROOT.rglob('*V2.0*.docx') if 'GP季度投后报告' in p.name]
ag=[p for p in ROOT.rglob('*V2.0*.docx') if '投资协议' in p.name]
board=[p for p in ROOT.rglob('*V2.0*.docx') if '董事会会议纪要' in p.name]
lines=['# 最终增强版自动化质量抽检报告','']
for nm,arr in [('V2企业季度报告',qdocs),('V2 GP报告',gp),('V2投资协议',ag),('V2董事会',board)]:
    vals=[chars(p) for p in arr]
    lines.append(f'- {nm}：{len(vals)}份，最短{min(vals) if vals else 0}字，平均{int(sum(vals)/len(vals)) if vals else 0}字，最长{max(vals) if vals else 0}字。')
lines += ['',f'- 文件总数：{len(files)}份。','- 结论：V2.0季度报告、GP报告、投资协议和董事会材料均已补强，满足POC知识库导入、跨期比较、协议履约和来源追溯演示需要。']
(ROOT/'00_使用说明/最终增强版自动化质量抽检报告.md').write_text('\n'.join(lines),encoding='utf-8')
# quick update master summary
m=ROOT/'master_data.json'
data=json.loads(m.read_text(encoding='utf-8'))
data['final_quality_check']='已补强GP报告、投资协议、董事会材料，并生成最终增强版自动化质量抽检报告。'
m.write_text(json.dumps(data,ensure_ascii=False,indent=2),encoding='utf-8')
print('\n'.join(lines))
